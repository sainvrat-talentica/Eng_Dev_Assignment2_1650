#!/usr/bin/env python3
"""
Refresh Assignment 2 sample use case outputs from live analytics API.

1. POST /api/v1/admin/analytics/import
2. Call UC1–UC6 endpoints
3. Write deliverables/raw-responses/uc*.json
4. Regenerate SAMPLE-USE-CASE-OUTPUTS.md and .docx

Usage:
  python3 scripts/refresh_analytics_outputs.py
  BASE_URL=http://localhost:8080 ADMIN_API_KEY=... python3 scripts/refresh_analytics_outputs.py
"""
from __future__ import annotations

import json
import os
import sys
import urllib.error
import urllib.request
from datetime import datetime, timezone
from pathlib import Path

A2_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = A2_ROOT / "deliverables"
RAW_DIR = OUT_DIR / "raw-responses"
MD_OUT = OUT_DIR / "SAMPLE-USE-CASE-OUTPUTS.md"
DOCX_OUT = OUT_DIR / "SAMPLE-USE-CASE-OUTPUTS.docx"

BASE_URL = os.environ.get("BASE_URL", "http://localhost:8080").rstrip("/")
ADMIN_KEY = os.environ.get("ADMIN_API_KEY", "local-dev-admin-key-change-me")

USE_CASES = [
    {
        "id": "uc1",
        "title": "UC1: Why were deliveries delayed in city X yesterday?",
        "question": "Why were deliveries delayed in Pune on 2025-06-06?",
        "method": "GET",
        "path": "/api/v1/analytics/delays?city=Pune&date=2025-06-06",
        "body": None,
    },
    {
        "id": "uc2",
        "title": "UC2: Why did Client X's orders fail in the past week?",
        "question": "Why did client 337 (Ramesh-Choudhary) orders fail in 2025?",
        "method": "GET",
        "path": "/api/v1/analytics/failures?clientId=337&from=2025-01-01T00:00:00Z&to=2025-12-31T00:00:00Z",
        "body": None,
    },
    {
        "id": "uc3",
        "title": "UC3: Top reasons for failures linked to Warehouse B in August",
        "question": "Explain the top reasons for delivery failures linked to Warehouse 2 in August 2025",
        "method": "GET",
        "path": "/api/v1/analytics/failures/by-warehouse?warehouseId=2&monthParam=2025-08",
        "body": None,
    },
    {
        "id": "uc4",
        "title": "UC4: Compare delivery failure causes between City A and City B",
        "question": "Compare delivery failure causes between Pune and Mumbai in August 2025",
        "method": "GET",
        "path": "/api/v1/analytics/failures/compare?cityA=Pune&cityB=Mumbai&monthParam=2025-08",
        "body": None,
    },
    {
        "id": "uc5",
        "title": "UC5: Festival period failures and preparation",
        "question": "What caused delivery failures during festival/holiday periods (Aug 2025)?",
        "method": "POST",
        "path": "/api/v1/analytics/insights/query",
        "body": {
            "queryType": "FESTIVAL_ANALYSIS",
            "parameters": {
                "from": "2025-08-01T00:00:00Z",
                "to": "2025-09-01T00:00:00Z",
            },
        },
    },
    {
        "id": "uc6",
        "title": "UC6: Onboard Client Y (+20,000 orders/month) — capacity risk",
        "question": "If client 337 adds 20,000 monthly orders, what failure risks emerge?",
        "method": "GET",
        "path": "/api/v1/analytics/capacity-projection?clientId=337&additionalMonthlyOrders=20000",
        "body": None,
    },
]


def http_request(method: str, path: str, body: dict | None = None, admin: bool = False) -> dict:
    url = f"{BASE_URL}{path}"
    data = json.dumps(body).encode("utf-8") if body is not None else None
    headers = {"Accept": "application/json"}
    if body is not None:
        headers["Content-Type"] = "application/json"
    if admin:
        headers["X-Admin-Api-Key"] = ADMIN_KEY
    req = urllib.request.Request(url, data=data, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            raw = resp.read().decode("utf-8")
            return json.loads(raw) if raw.strip() else {}
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"{method} {path} failed HTTP {e.code}: {detail}") from e


def import_dataset() -> dict:
    print("Importing analytics CSV dataset...")
    return http_request("POST", "/api/v1/admin/analytics/import", admin=True)


def format_md_section(uc: dict, response: dict) -> str:
    api_line = f"`{uc['method']} {uc['path']}`"
    if uc["body"]:
        api_line = f"`{uc['method']} {uc['path']}`"
    lines = [
        f"## {uc['title']}",
        "",
        f"**Business question:** {uc['question']}",
        "",
        f"**API:** {api_line}",
        "",
    ]
    if uc["body"]:
        lines.extend([
            "**Request body:**",
            "",
            "```json",
            json.dumps(uc["body"], indent=2),
            "```",
            "",
        ])
    narrative = response.get("narrative")
    if narrative:
        lines.extend([
            "### Narrative (human-readable insight)",
            "",
            narrative,
            "",
        ])
    recommendations = response.get("recommendations")
    if recommendations:
        lines.extend(["### Recommendations", ""])
        for rec in recommendations:
            lines.append(f"- {rec}")
        lines.append("")
    lines.extend([
        "### Evidence / structured output",
        "",
        "```json",
        json.dumps(response, indent=2),
        "```",
        "",
        "---",
        "",
    ])
    return "\n".join(lines)


def write_docx(md_text: str, path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("Skip .docx (install python-docx)", file=sys.stderr)
        return
    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], 0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], 2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("```"):
            continue
        elif line.strip() == "---":
            continue
        elif line.strip():
            p = doc.add_paragraph(line)
            for r in p.runs:
                r.font.size = Pt(10)
    doc.save(path)


def main() -> int:
    RAW_DIR.mkdir(parents=True, exist_ok=True)

    # Health check
    try:
        health = http_request("GET", "/actuator/health")
        if health.get("status") != "UP":
            print(f"Warning: health status = {health.get('status')}", file=sys.stderr)
    except Exception as e:
        print(f"API not reachable at {BASE_URL}: {e}", file=sys.stderr)
        return 1

    import_result = import_dataset()
    print(f"Import: success={import_result.get('success')} rows={import_result.get('rowCounts', {})}")

    sections = []
    for uc in USE_CASES:
        print(f"Fetching {uc['id']}...")
        if uc["method"] == "GET":
            response = http_request("GET", uc["path"])
        else:
            response = http_request("POST", uc["path"], body=uc["body"])
        raw_path = RAW_DIR / f"{uc['id']}.json"
        raw_path.write_text(json.dumps(response, indent=4) + "\n", encoding="utf-8")
        print(f"  Wrote {raw_path}")
        sections.append(format_md_section(uc, response))

    recorded = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    header = f"""# Assignment 2 — Recorded Sample Use Case Outputs

**Project:** SwiftEats — Delivery Failure Root-Cause Analytics  
**Recorded:** {recorded}  
**Environment:** Docker Compose + SwiftEats API at `{BASE_URL}`  
**Dataset:** `sample-data/` imported into PostgreSQL schema `analytics.*`  
**Import result:** success={import_result.get('success')} — see row counts in import API response below

## How to reproduce

```bash
docker compose up -d
curl -X POST "{BASE_URL}/api/v1/admin/analytics/import" \\
  -H "X-Admin-Api-Key: {ADMIN_KEY}"

python3 scripts/refresh_analytics_outputs.py
```

Or refresh manually with the API calls listed under each use case below.

Raw JSON responses: [`raw-responses/`](./raw-responses/)

### Import response (this run)

```json
{json.dumps(import_result, indent=2)}
```

---

"""
    md_content = header + "\n".join(sections)
    MD_OUT.write_text(md_content, encoding="utf-8")
    print(f"Wrote {MD_OUT}")
    write_docx(md_content, DOCX_OUT)
    if DOCX_OUT.exists():
        print(f"Wrote {DOCX_OUT}")
    print("Done — all 6 use cases refreshed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

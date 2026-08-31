#!/usr/bin/env python3
"""Generate deliverables/ASSIGNMENT-2-WRITEUP.docx from project content."""

import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("Install python-docx: pip install python-docx")

A2_ROOT = Path(__file__).resolve().parents[1]
OUT = A2_ROOT / "deliverables" / "ASSIGNMENT-2-WRITEUP.docx"
DIAGRAM = A2_ROOT / "deliverables" / "architecture-diagram.png"
DIAGRAM_SCRIPT = A2_ROOT / "scripts" / "generate_architecture_diagram.py"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def ensure_diagram() -> Path:
    if not DIAGRAM.exists() or DIAGRAM.stat().st_mtime < DIAGRAM_SCRIPT.stat().st_mtime:
        subprocess.run([sys.executable, str(DIAGRAM_SCRIPT)], check=True)
    return DIAGRAM


def main():
    doc = Document()
    title = doc.add_heading("Assignment 2 — Delivery Failure Root-Cause Analytics", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "SwiftEats Platform — Problem Solution Write-Up", bold=True)
    add_para(doc, "Author: SwiftEats Team | Date: August 2026")
    doc.add_paragraph()

    add_heading(doc, "1. Problem Statement", 1)
    add_para(
        doc,
        "Delivery failures and delays drive customer dissatisfaction and revenue leakage. "
        "Operations teams can count failures but struggle to explain why they happened because data is fragmented:"
    )
    add_bullets(
        doc,
        [
            "Order & shipment timestamps exist but are rarely linked to external conditions (traffic, weather).",
            "Fleet GPS traces and driver notes are unstructured and not analyzed systematically.",
            "Warehouse logs (stockouts, prep delays) are not tied to downstream delivery outcomes.",
            "Customer feedback is free-text and hard to aggregate.",
            "Contextual data (traffic, weather, festivals) sits in separate systems.",
        ],
    )
    add_para(
        doc,
        "The strategic need is a system that aggregates multi-domain data, correlates events automatically, "
        "generates human-readable insights, and surfaces actionable recommendations."
    )

    add_heading(doc, "2. Proposed Solution", 1)
    add_para(
        doc,
        "We implemented an Analytics & Insights module as part of the SwiftEats platform. "
        "It ingests the provided CSV sample dataset into a dedicated analytics schema, "
        "enriches orders with warehouse, fleet, feedback, and external-factor context, "
        "runs a rule-based correlation engine, and produces narrative insights plus recommendations."
    )
    add_para(doc, "The sample program exposes:", bold=False)
    add_bullets(
        doc,
        [
            "REST query APIs via analytics-service on http://localhost:8080.",
            "Optional CLI demo runner (AnalyticsDemoRunner) for video walkthroughs.",
            "Optional analytics dashboard UI (http://localhost:3002) for interactive demos.",
        ],
    )
    add_para(
        doc,
        "This exceeds the assignment minimum (a simple local program) while keeping the core deliverable "
        "a focused aggregation + correlation + narrative pipeline."
    )

    add_heading(doc, "3. Architecture Overview", 1)
    add_para(doc, "High-level data flow:", bold=True)
    diagram_path = ensure_diagram()
    doc.add_picture(str(diagram_path), width=Inches(6.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Figure 1 — Multi-domain CSV import, correlation, and insight generation pipeline."
    )

    add_para(doc, "Component responsibilities:", bold=True)
    add_bullets(
        doc,
        [
            "Import layer: Parses CSV, normalizes fields, loads analytics tables with import lock.",
            "Query layer: AnalyticsQueryService implements UC1–UC6 with typed SQL aggregations.",
            "CorrelationEngine: Matches failure reasons with warehouse notes, fleet notes, traffic, weather, festival events.",
            "InsightGenerator: Converts structured results into operations-friendly paragraphs and recommendation lists.",
        ],
    )

    add_heading(doc, "4. Data Sources", 1)
    add_para(doc, "Sample dataset: sample-data/", bold=True)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "File"
    hdr[1].text = "Purpose"
    rows = [
        ("orders.csv", "Order timestamps, failure/delay flags, client and warehouse linkage"),
        ("clients.csv", "B2B client master for per-client failure analysis (UC2, UC6)"),
        ("warehouses.csv", "Warehouse capacity and location for UC3 and capacity projection"),
        ("warehouse_logs.csv", "Prep delays, stockouts, system issues"),
        ("fleet_logs.csv", "GPS delay notes, congestion, address issues"),
        ("feedback.csv", "Customer complaints linked to orders"),
        ("external_factors.csv", "Traffic, weather, festival/holiday context"),
        ("drivers.csv", "Fleet reference data"),
    ]
    for f, p in rows:
        row = table.add_row().cells
        row[0].text = f
        row[1].text = p

    add_heading(doc, "5. Correlation Logic", 1)
    add_para(
        doc,
        "The CorrelationEngine applies deterministic rules (no paid LLM required) to enriched order rows:"
    )
    add_bullets(
        doc,
        [
            "STOCKOUT_WAREHOUSE: failure reason Stockout + warehouse note mentions stock delay.",
            "WAREHOUSE_OPS: Warehouse delay + slow packing or system issue notes.",
            "TRAFFIC_TRIPLE_CONFIRM: Traffic congestion failure + heavy fleet congestion + heavy traffic condition.",
            "ADDRESS_MISMATCH: Incorrect address failure + address-not-found fleet note.",
            "WEATHER_IMPACT: Weather disruption + rain/fog conditions.",
            "FESTIVAL_VOLUME: External event type Festival or Holiday during failure window.",
            "SLA_BREACH: Delayed but not failed orders.",
        ],
    )
    add_para(
        doc,
        "InsightGenerator maps matched rules to template narratives and operational recommendations "
        "(e.g., pre-scale warehouse capacity, add driver buffer before festivals, enable inventory alerts)."
    )

    add_heading(doc, "6. Sample Use Case Walkthrough", 1)
    use_cases = [
        (
            "UC1 — Delays in city X",
            "GET /api/v1/analytics/delays?city=Pune&date=2025-06-06",
            "Aggregates affected orders in Pune on the date; ranks failure reasons; counts traffic, slow packing, and negative feedback correlations.",
        ),
        (
            "UC2 — Client X failures",
            "GET /api/v1/analytics/failures?clientId=337&from=...&to=...",
            "Groups failed orders for client 337 with breakdown by failure reason, fleet, and warehouse notes.",
        ),
        (
            "UC3 — Warehouse B failures in August",
            "GET /api/v1/analytics/failures/by-warehouse?warehouseId=2&month=2025-08",
            "Top failure reasons and correlated operational signals for a warehouse in a month.",
        ),
        (
            "UC4 — Compare City A vs City B",
            "GET /api/v1/analytics/failures/compare?cityA=Pune&cityB=Mumbai&month=2025-08",
            "Side-by-side failure reason distribution and volume comparison.",
        ),
        (
            "UC5 — Festival period preparation",
            "POST /api/v1/analytics/insights/query with queryType FESTIVAL_ANALYSIS",
            "Analyzes failures during festival/holiday windows; recommends capacity buffers.",
        ),
        (
            "UC6 — Client Y capacity risk (+20K orders/month)",
            "GET /api/v1/analytics/capacity-projection?clientId=337&additionalMonthlyOrders=20000",
            "Projects warehouse utilization and flags high-risk sites.",
        ),
    ]
    for title, api, desc in use_cases:
        add_heading(doc, title, 2)
        add_para(doc, f"API: {api}")
        add_para(doc, desc)
    add_para(
        doc,
        "Full recorded outputs with narratives and JSON evidence: "
        "deliverables/SAMPLE-USE-CASE-OUTPUTS.md"
    )

    add_heading(doc, "7. How to Run the Demo Locally", 1)
    add_bullets(
        doc,
        [
            "docker compose up --build",
            "Import data: POST /api/v1/admin/analytics/import with X-Admin-Api-Key",
            "Run queries above via curl, Postman, analytics dashboard (:3002), or AnalyticsDemoRunner",
            "See README.md and deliverables/DEMO-SCRIPT.md for video recording steps",
        ],
    )

    add_heading(doc, "8. Design Trade-offs", 1)
    add_bullets(
        doc,
        [
            "Rule-based correlation chosen over LLM for determinism, testability, and offline demo reliability.",
            "Batch CSV import matches assignment sample data; live Kafka correlation is a future evolution.",
            "Restaurant entity maps to warehouse in sample queries — documented in DOMAIN_MODEL.md.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
